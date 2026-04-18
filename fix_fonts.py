#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""修复DOCX文件中的字体问题，强制使用正确的中英文字体"""

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import sys
from pathlib import Path


def fix_fonts_in_docx(docx_path):
    """
    强制修复docx中的字体和颜色：
    - 标题：黑体(SimHei) + 粗体 + 黑色
    - 正文：宋体(SimSun)
    - 英文：Times New Roman
    - 清除所有颜色和斜体
    """
    from docx.oxml import OxmlElement
    from copy import deepcopy

    doc = Document(docx_path)

    def clean_run(run, is_heading):
        run.font.name = 'Times New Roman'
        if is_heading:
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
            run.font.bold = True
        else:
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
            run.font.bold = False
        run.font.italic = False
        rPr = run._element.rPr
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            run._element.insert(0, rPr)
        for color in rPr.findall(qn('w:color')):
            rPr.remove(color)
        for shd in rPr.findall(qn('w:shd')):
            rPr.remove(shd)

    # Fix all runs in paragraphs
    for paragraph in doc.paragraphs:
        is_heading = paragraph.style and paragraph.style.name.startswith('Heading')
        for run in paragraph.runs:
            clean_run(run, is_heading)

    # Fix all runs in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    is_heading = paragraph.style and paragraph.style.name.startswith('Heading')
                    for run in paragraph.runs:
                        clean_run(run, is_heading)

    # Fix Heading style definitions: clear color and italic from style XML
    heading_style_names = {f'Heading {i}' for i in range(1, 10)}
    styles_to_fix = [s for s in doc.styles if s.name in heading_style_names]

    for style in styles_to_fix:
        style_elem = style._element
        # Remove color from style definition (depth-first search in XML tree)
        for elem in [style_elem] + list(style_elem.iter()):
            for color in list(elem.findall(qn('w:color'))):
                elem.remove(color)
            for shd in list(elem.findall(qn('w:shd'))):
                elem.remove(shd)
            for tlc in list(elem.findall(qn('w:themeColor'))):
                elem.remove(tlc)

        # Ensure bold and roman fonts in style
        rPr = style_elem.find(qn('w:rPr'))
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            style_elem.append(rPr)
        # Set font to Times New Roman / SimHei
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rFonts.set(qn('w:eastAsia'), 'SimHei')
        rFonts.set(qn('w:cs'), 'Times New Roman')
        # Bold
        b = rPr.find(qn('w:b'))
        if b is None:
            b = OxmlElement('w:b')
            rPr.append(b)
        # No italic
        for i in rPr.findall(qn('w:i')):
            rPr.remove(i)

    doc.save(docx_path)
    print(f"[OK] Fixed: {docx_path}")
    return True


def main():
    if len(sys.argv) < 2:
        print("Usage: python fix_fonts.py <file.docx> [file2.docx ...]")
        print("   or: python fix_fonts.py --batch <directory>")
        sys.exit(1)

    if sys.argv[1] == '--batch' and len(sys.argv) >= 3:
        # 批量处理目录
        dir_path = Path(sys.argv[2])
        docx_files = list(dir_path.glob('*.docx'))
        print(f"[SCAN] Found {len(docx_files)} docx files in {dir_path}")
        for f in docx_files:
            fix_fonts_in_docx(f)
    else:
        # 处理单个或多个文件
        for path in sys.argv[1:]:
            fix_fonts_in_docx(path)


if __name__ == '__main__':
    main()
