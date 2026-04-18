#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""创建带字体设置的 Word 模板"""

from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn


def add_heading_styled(doc, text, level=1):
    """添加标题并设置黑体字体"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')  # 黑体
        if level == 1:
            run.font.size = Pt(16)
            run.font.bold = True
        elif level == 2:
            run.font.size = Pt(14)
            run.font.bold = True
        else:
            run.font.size = Pt(12)
    return heading


def add_paragraph_styled(doc, text, font_size=12):
    """添加正文段落并设置宋体字体"""
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')  # 宋体
        run.font.size = Pt(font_size)
    return p


def create_template(output_path):
    """创建模板文档"""
    doc = Document()

    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    # 添加标题示例
    add_heading_styled(doc, 'Chapter 1 Template Title', level=1)
    add_heading_styled(doc, '1.1 Section Title', level=2)

    # 添加正文
    add_paragraph_styled(doc, 'This is body text. This template sets:')
    add_paragraph_styled(doc, '- Heading font: SimHei ( 黑体 )')
    add_paragraph_styled(doc, '- Body font: SimSun ( 宋体 )')
    add_paragraph_styled(doc, '- English font: Times New Roman')
    add_paragraph_styled(doc, '- Heading size: 16pt (三号)')
    add_paragraph_styled(doc, '- Body size: 12pt (小四)')

    # 添加英文示例
    doc.add_paragraph()
    p = doc.add_paragraph('English text: The quick brown fox jumps over the lazy dog.')
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

    # 添加中文示例
    doc.add_paragraph()
    add_paragraph_styled(doc, '中文示例：本文档使用黑体作为标题字体，宋体作为正文字体。')

    # 保存模板
    doc.save(output_path)
    print(f'Template created: {output_path}')


if __name__ == '__main__':
    output_path = 'md2docx-tool/templates/default.docx'
    create_template(output_path)
