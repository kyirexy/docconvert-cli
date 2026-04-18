# -*- coding: utf-8 -*-
"""字体修复模块 - 强制设置中文/英文字体，清除颜色和斜体"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def fix_fonts_in_docx(docx_path) -> bool:
    """
    修复 docx 中的字体和颜色：
    - 标题：黑体(SimHei) + 粗体 + 黑色
    - 正文：宋体(SimSun)
    - 英文：Times New Roman
    - 清除所有颜色和斜体
    """
    def clean_run(run, is_heading):
        run.font.name = 'Times New Roman'
        if is_heading:
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
            run.font.bold = True
        else:
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
            run.font.bold = False
        run.font.italic = False
        rPr = run._element.rPr
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            run._element.insert(0, rPr)
        for color in rPr.findall(qn('w:color')):
            rPr.remove(color)
        for shd in rPr.findall(qn('w:shd')):
            rPr.remove(shd)

    try:
        doc = Document(docx_path)

        # 修复所有段落中的 run
        for paragraph in doc.paragraphs:
            is_heading = paragraph.style and paragraph.style.name.startswith('Heading')
            for run in paragraph.runs:
                clean_run(run, is_heading)

        # 修复所有表格中的 run
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        is_heading = paragraph.style and paragraph.style.name.startswith('Heading')
                        for run in paragraph.runs:
                            clean_run(run, is_heading)

        # 修复 Heading 样式定义
        heading_style_names = {f'Heading {i}' for i in range(1, 10)}
        for style in doc.styles:
            if style.name not in heading_style_names:
                continue
            style_elem = style._element
            for elem in [style_elem] + list(style_elem.iter()):
                for color in list(elem.findall(qn('w:color'))):
                    elem.remove(color)
                for shd in list(elem.findall(qn('w:shd'))):
                    elem.remove(shd)
                for tlc in list(elem.findall(qn('w:themeColor'))):
                    elem.remove(tlc)
            # 设置字体和粗体
            rPr = style_elem.find(qn('w:rPr'))
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                style_elem.append(rPr)
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.append(rFonts)
            rFonts.set(qn('w:ascii'), 'Times New Roman')
            rFonts.set(qn('w:hAnsi'), 'Times New Roman')
            rFonts.set(qn('w:eastAsia'), 'SimHei')
            rFonts.set(qn('w:cs'), 'Times New Roman')
            if rPr.find(qn('w:b')) is None:
                rPr.append(OxmlElement('w:b'))
            for i in rPr.findall(qn('w:i')):
                rPr.remove(i)

        doc.save(docx_path)
        return True
    except Exception:
        return False
